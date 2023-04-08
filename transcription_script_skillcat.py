import os

os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "C:\Json Keys\skillcat-transcription-script-475f3a0747da.json"
import docx
from google.cloud import translate_v2 as translate

# Instantiates a client
translate_client = translate.Client()

# Read the document
doc = docx.Document('1.docx')

# Loop through each table in the document
for table in doc.tables:
    if len(table.columns) < 4:  # Skip tables with less than 4 columns
        continue
    for row in table.rows:
        cell = row.cells[3]  # Only translate the text in the fourth column
        # Ignore text elements that should not be translated
        if cell.text.strip() == "Translation" or "%Project.SlideNumber%/%Project.TotalSlides%" in cell.text:
            continue
        # Translate the text in the cell
        result = translate_client.translate(cell.text, target_language='es')
        # Replace the text in the cell with the translated text
        cell.text = result['translatedText']
        # Preserve the formatting of the original text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if run.text == result['input']:
                    run.text = result['translatedText']

# Save the updated document
doc.save('updated_document.docx')
